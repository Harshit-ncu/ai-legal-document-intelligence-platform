# app/services/extractor.py
# ─────────────────────────────────────────────────────────
# Text extraction service for PDF, DOCX, and TXT files.
#
# DESIGN RULES for this file:
#   - No HTTP code (no Request, Response, FastAPI imports)
#   - Each function takes a file PATH (str) and returns a dict
#   - Raises ValueError for known bad inputs (wrong format)
#   - Raises RuntimeError for unexpected parsing failures
#
# This separation means you can call these functions directly
# in a test or a script — no web server needed.
# ─────────────────────────────────────────────────────────

import fitz          # PyMuPDF — 'fitz' is its internal module name
import docx          # python-docx
from pathlib import Path


# ── Helper ────────────────────────────────────────────────

def _count_words(text: str) -> int:
    """
    Split text on whitespace and count non-empty tokens.
    'split()' with no argument handles any whitespace
    (spaces, newlines, tabs) and ignores leading/trailing space.
    """
    return len(text.split())


# ── PDF Extractor ─────────────────────────────────────────

def extract_pdf_text(file_path: str) -> dict:
    """
    Extract text from a PDF file using PyMuPDF.

    How PyMuPDF works:
      1. fitz.open() loads the PDF into memory.
      2. We iterate over every Page object.
      3. page.get_text("text") extracts all text blocks on that
         page as a single string, preserving reading order.
      4. Pages are joined with double newlines so the reader
         can tell where one page ends and the next begins.

    Why PyMuPDF over PyPDF2?
      PyMuPDF uses MuPDF — the same engine used by Firefox and
      Sumatra PDF. It handles complex layouts, embedded fonts,
      and multi-column text far better than PyPDF2, which is
      now archived and unmaintained.

    Args:
        file_path: Absolute path to the .pdf file on disk.

    Returns:
        dict with keys: text, pageCount, wordCount, documentType
    """
    try:
        # fitz.open() supports PDF, XPS, EPUB, and more.
        # The context manager ensures the document is closed
        # and memory is freed even if an error occurs.
        doc = fitz.open(file_path)
    except Exception as e:
        raise ValueError(f"Cannot open PDF file: {e}")

    with doc:
        page_count = len(doc)  # total number of pages

        if page_count == 0:
            raise ValueError("PDF has no pages.")

        page_texts = []

        for page_num in range(page_count):
            page = doc[page_num]

            # "text" mode → plain text in reading order.
            # Other modes: "html", "dict", "blocks", "words"
            # We use plain "text" — clean and simple.
            page_text = page.get_text("text").strip()

            # Only append if the page actually has text.
            # Some PDFs have blank pages or image-only pages.
            if page_text:
                page_texts.append(page_text)

        # Join pages with double newline — readable separator
        full_text = "\n\n".join(page_texts)

    return {
        "documentType": "pdf",
        "pageCount":    page_count,
        "wordCount":    _count_words(full_text),
        "text":         full_text,
    }


# ── DOCX Extractor ────────────────────────────────────────

def extract_docx_text(file_path: str) -> dict:
    """
    Extract text from a .docx (Word) file using python-docx.

    How python-docx works:
      A .docx file is actually a ZIP archive containing XML files.
      python-docx unzips it and parses the XML automatically.
      The main content lives in word/document.xml.

      The Document object exposes content as:
        - doc.paragraphs   → list of Paragraph objects
        - paragraph.text   → the plain text of that paragraph

      Note: Tables, headers, and footers are NOT in doc.paragraphs.
      We add table extraction below for completeness.

    What we report as "pageCount":
      Word documents don't have a fixed page count at the XML level
      (page breaks depend on the render engine + font + printer).
      We count paragraphs instead — a meaningful unit of content.

    Args:
        file_path: Absolute path to the .docx file on disk.

    Returns:
        dict with keys: text, pageCount, wordCount, documentType
    """
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        raise ValueError(f"Cannot open DOCX file: {e}")

    # ── Extract from paragraphs ───────────────────────────
    # doc.paragraphs is a list of Paragraph objects.
    # Each paragraph.text gives the string content.
    # We filter empty strings (blank lines between sections).
    paragraph_texts = [
        para.text.strip()
        for para in doc.paragraphs
        if para.text.strip()  # skip blank paragraphs
    ]

    # ── Extract from tables ───────────────────────────────
    # Tables in Word docs (e.g. contract schedules, clauses)
    # are NOT included in doc.paragraphs.
    # We iterate each table → each row → each cell.
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            cell_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cell_texts:
                # Join cells in a row with a tab character for readability
                table_texts.append("\t".join(cell_texts))

    # Combine: main body paragraphs first, then table content
    all_parts = paragraph_texts + table_texts
    full_text = "\n\n".join(all_parts)

    return {
        "documentType": "docx",
        # Report paragraph count as the meaningful unit
        "pageCount":    len(paragraph_texts),
        "wordCount":    _count_words(full_text),
        "text":         full_text,
    }


# ── TXT Extractor ─────────────────────────────────────────

def extract_txt_text(file_path: str) -> dict:
    """
    Extract text from a plain .txt file.

    How it works:
      Standard Python file I/O with explicit UTF-8 encoding.
      'errors="replace"' means malformed bytes are replaced
      with the Unicode replacement character (U+FFFD) instead
      of raising a UnicodeDecodeError — handles legacy files
      gracefully.

    What we report as "pageCount":
      TXT files have no pages. We count non-empty lines instead
      — a natural unit for plain text documents.

    Args:
        file_path: Absolute path to the .txt file on disk.

    Returns:
        dict with keys: text, pageCount, wordCount, documentType
    """
    try:
        path = Path(file_path)

        # Read the entire file as a UTF-8 string.
        # 'errors="replace"' avoids crashes on non-UTF-8 bytes.
        full_text = path.read_text(encoding="utf-8", errors="replace")
    except FileNotFoundError:
        raise ValueError(f"File not found: {file_path}")
    except Exception as e:
        raise RuntimeError(f"Cannot read TXT file: {e}")

    # Count non-empty lines as the "page" equivalent
    non_empty_lines = [line for line in full_text.splitlines() if line.strip()]
    line_count = len(non_empty_lines)

    # Clean up extra blank lines (collapse 3+ newlines to 2)
    import re
    clean_text = re.sub(r"\n{3,}", "\n\n", full_text).strip()

    return {
        "documentType": "txt",
        "pageCount":    line_count,
        "wordCount":    _count_words(clean_text),
        "text":         clean_text,
    }
